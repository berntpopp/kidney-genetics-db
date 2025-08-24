"""DOCX file extractor."""

import logging
from pathlib import Path
from typing import Any, Dict, List

from docx import Document


class DocxExtractor:
    """Extract content from DOCX files."""

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

    def extract(self, file_path: Path) -> Dict[str, Any]:
        """Extract all content from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted content
        """
        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "success": True,
            }

        except Exception as e:
            self.logger.error(f"Error extracting DOCX {file_path}: {e}")
            return {"success": False, "error": str(e)}

    def extract_tables_with_metadata(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract tables with cell metadata for precise extraction.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of table dictionaries with cell metadata
        """
        try:
            doc = Document(file_path)
            tables_with_metadata = []

            for table_idx, table in enumerate(doc.tables):
                table_cells = []
                cell_id = 0

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_id += 1
                        table_cells.append(
                            {
                                "cell_id": cell_id,
                                "row": row_idx,
                                "col": col_idx,
                                "text": cell.text.strip(),
                            }
                        )

                tables_with_metadata.append(
                    {
                        "table_index": table_idx,
                        "cells": table_cells,
                    }
                )

            return tables_with_metadata

        except Exception as e:
            self.logger.error(f"Error extracting tables from DOCX {file_path}: {e}")
            return []
